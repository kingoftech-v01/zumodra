"""
Resume Parser Service

Extracts structured data from PDF and DOCX resumes including:
- Contact information (name, email, phone)
- Skills and competencies
- Work experience
- Education
- Certifications
"""
import io
import logging
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


@dataclass
class ParsedResume:
    """Structured data extracted from a resume."""
    name: str = ''
    email: str = ''
    phone: str = ''
    location: str = ''
    summary: str = ''
    skills: List[str] = field(default_factory=list)
    experience: List[Dict[str, Any]] = field(default_factory=list)
    education: List[Dict[str, Any]] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    languages: List[str] = field(default_factory=list)
    total_experience_years: float = 0.0
    raw_text: str = ''


class ResumeParserService:
    """
    Service for parsing and extracting data from resumes.

    Supports PDF and DOCX formats with fallback parsers.

    Usage:
        service = ResumeParserService()
        with open('resume.pdf', 'rb') as f:
            result = service.parse(f.read(), 'pdf')
        print(f"Skills: {result['skills']}")
    """

    # Common technical skill patterns
    SKILL_PATTERNS = [
        # Programming languages
        r'\b(python|java|javascript|typescript|c\+\+|c#|ruby|go|golang|rust|'
        r'scala|kotlin|swift|objective-c|php|perl|r\b|matlab|julia)\b',
        # Web frameworks
        r'\b(react|angular|vue|svelte|django|flask|fastapi|spring|node\.?js|'
        r'express|rails|laravel|asp\.net|next\.?js|nuxt)\b',
        # Cloud & DevOps
        r'\b(aws|azure|gcp|google cloud|docker|kubernetes|k8s|terraform|'
        r'ansible|jenkins|gitlab|github actions|ci/cd|devops)\b',
        # Databases
        r'\b(postgresql|postgres|mysql|mongodb|redis|elasticsearch|'
        r'dynamodb|cassandra|sqlite|oracle|sql server|mariadb)\b',
        # Data & ML
        r'\b(machine learning|deep learning|nlp|natural language|'
        r'computer vision|tensorflow|pytorch|keras|scikit-learn|pandas|numpy)\b',
        # Other technologies
        r'\b(git|linux|unix|bash|shell|rest api|graphql|grpc|microservices|'
        r'agile|scrum|jira|confluence)\b',
    ]

    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    PHONE_PATTERN = (
        r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?'
        r'[-\s\.]?[0-9]{3,6}[-\s\.]?[0-9]{3,6}'
    )

    # Section headers to look for
    SECTION_PATTERNS = {
        'experience': [
            r'experience', r'work history', r'employment', r'professional experience',
            r'work experience', r'career history'
        ],
        'education': [
            r'education', r'academic', r'qualifications', r'degrees', r'schooling'
        ],
        'skills': [
            r'skills', r'technical skills', r'competencies', r'expertise',
            r'technologies', r'proficiencies'
        ],
        'summary': [
            r'summary', r'profile', r'objective', r'about', r'overview'
        ],
        'certifications': [
            r'certifications?', r'certificates?', r'licenses?', r'credentials'
        ],
    }

    def __init__(self):
        self._compiled_patterns = None

    def parse(
        self,
        file_content: bytes,
        file_type: str = 'pdf'
    ) -> Dict[str, Any]:
        """
        Parse a resume file and extract structured data.

        Args:
            file_content: Binary content of the file
            file_type: 'pdf' or 'docx'

        Returns:
            Dict with extracted resume data
        """
        try:
            # Extract raw text based on file type
            file_type = file_type.lower()
            if file_type == 'pdf':
                raw_text = self._extract_pdf_text(file_content)
            elif file_type in ('docx', 'doc'):
                raw_text = self._extract_docx_text(file_content)
            elif file_type == 'txt':
                raw_text = file_content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Parse the text
            parsed = self._parse_text(raw_text)

            return {
                'name': parsed.name,
                'email': parsed.email,
                'phone': parsed.phone,
                'location': parsed.location,
                'summary': parsed.summary,
                'skills': parsed.skills,
                'experience': parsed.experience,
                'education': parsed.education,
                'certifications': parsed.certifications,
                'languages': parsed.languages,
                'total_experience_years': parsed.total_experience_years,
                'raw_text': raw_text[:5000],  # Truncate for storage
            }

        except Exception as e:
            logger.error(f"Resume parsing failed: {e}", exc_info=True)
            return {
                'name': '',
                'email': '',
                'phone': '',
                'location': '',
                'summary': '',
                'skills': [],
                'experience': [],
                'education': [],
                'certifications': [],
                'languages': [],
                'total_experience_years': 0.0,
                'error': str(e)
            }

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content."""
        # Try pdfplumber first (better layout preservation)
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return '\n'.join(text_parts)

        except ImportError:
            logger.debug("pdfplumber not available, trying PyPDF2")

        # Fallback to PyPDF2
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return '\n'.join(text_parts)

        except ImportError:
            logger.error("No PDF parser available. Install pdfplumber or PyPDF2")
            raise ImportError(
                "PDF parsing requires pdfplumber or PyPDF2. "
                "Install with: pip install pdfplumber"
            )

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return '\n'.join(text_parts)

        except ImportError:
            logger.error("python-docx not available")
            raise ImportError(
                "DOCX parsing requires python-docx. "
                "Install with: pip install python-docx"
            )

    def _parse_text(self, text: str) -> ParsedResume:
        """Parse extracted text into structured data."""
        parsed = ParsedResume(raw_text=text)

        if not text:
            return parsed

        # Extract contact information
        parsed.email = self._extract_email(text)
        parsed.phone = self._extract_phone(text)
        parsed.name = self._extract_name(text)
        parsed.location = self._extract_location(text)

        # Extract sections
        parsed.summary = self._extract_section(text, 'summary')
        parsed.skills = self.extract_skills(text)
        parsed.experience = self._extract_experience(text)
        parsed.education = self._extract_education(text)
        parsed.certifications = self._extract_certifications(text)
        parsed.languages = self._extract_languages(text)

        # Calculate total experience
        parsed.total_experience_years = self._calculate_total_experience(
            parsed.experience
        )

        return parsed

    def _extract_email(self, text: str) -> str:
        """Extract email address from text."""
        match = re.search(self.EMAIL_PATTERN, text)
        return match.group(0) if match else ''

    def _extract_phone(self, text: str) -> str:
        """Extract phone number from text."""
        match = re.search(self.PHONE_PATTERN, text)
        if match:
            # Clean up the phone number
            phone = match.group(0)
            phone = re.sub(r'[^\d+]', '', phone)
            return phone
        return ''

    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume text."""
        lines = text.split('\n')

        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if not line or len(line) > 60:
                continue

            # Skip lines that look like headers or contact info
            if '@' in line or re.match(r'^[\d\+\-\(\)]+', line):
                continue

            # Skip common section headers
            if any(h in line.lower() for h in ['resume', 'cv', 'curriculum']):
                continue

            # Basic name pattern (2-4 words, with capitals)
            if re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){1,3}$', line):
                return line

            # More lenient pattern
            if re.match(r'^[A-Z][a-zA-Z\-\']+(\s+[A-Z][a-zA-Z\-\']+){1,3}$', line):
                return line

        return ''

    def _extract_location(self, text: str) -> str:
        """Extract location from resume text."""
        # Look for common location patterns
        location_patterns = [
            r'(?:location|address|based in|residing in)[:\s]+([^\n,]{5,50})',
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z]{2})\s*(?:\d{5})?',  # City, ST
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',  # City, Country
        ]

        for pattern in location_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0).strip()

        return ''

    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from text.

        Args:
            text: Text to analyze

        Returns:
            List of extracted skills (lowercase, deduplicated)
        """
        skills = set()
        text_lower = text.lower()

        # Match against skill patterns
        for pattern in self.SKILL_PATTERNS:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                skill = match.strip().lower()
                if len(skill) > 1:  # Skip single characters
                    skills.add(skill)

        # Also look for skill sections
        skill_section = self._extract_section(text, 'skills')
        if skill_section:
            # Extract comma, bullet, or newline separated items
            items = re.split(r'[,\n•\|\-]', skill_section)
            for item in items:
                item = item.strip()
                if item and 2 < len(item) < 50:
                    # Basic validation - should contain letters
                    if re.search(r'[a-zA-Z]', item):
                        skills.add(item.lower())

        return sorted(list(skills))

    def _extract_section(self, text: str, section_type: str) -> str:
        """Extract content from a specific section."""
        patterns = self.SECTION_PATTERNS.get(section_type, [])

        for header in patterns:
            # Look for section header followed by content
            pattern = rf'(?i)(?:^|\n)\s*{header}\s*[:\n](.+?)(?=\n\s*(?:[A-Z][A-Za-z]+\s*[:\n])|$)'
            match = re.search(pattern, text, re.DOTALL)
            if match:
                content = match.group(1).strip()
                if len(content) > 10:
                    return content[:2000]  # Limit section length

        return ''

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience from resume."""
        experiences = []

        exp_section = self._extract_section(text, 'experience')
        if not exp_section:
            return experiences

        # Split by common job entry patterns
        # Look for patterns like: "Job Title at Company" or "Company - Job Title"
        job_patterns = [
            r'([A-Z][^,\n]{5,50})\s*(?:at|@|\|)\s*([A-Z][^,\n]{2,50})',
            r'([A-Z][^,\n]{2,50})\s*[-–—]\s*([A-Z][^,\n]{5,50})',
        ]

        for pattern in job_patterns:
            matches = re.findall(pattern, exp_section)
            for match in matches[:10]:  # Limit to 10 jobs
                title, company = match
                experiences.append({
                    'job_title': title.strip(),
                    'company': company.strip(),
                    'start_date': '',
                    'end_date': '',
                    'description': '',
                    'current': False
                })

        # If no matches found, try simpler extraction
        if not experiences:
            lines = exp_section.split('\n')
            for line in lines:
                line = line.strip()
                if line and 10 < len(line) < 100:
                    # Skip lines that look like descriptions
                    if not line.startswith(('•', '-', '*', '–')):
                        experiences.append({
                            'job_title': line,
                            'company': '',
                            'start_date': '',
                            'end_date': '',
                            'description': '',
                            'current': False
                        })
                        if len(experiences) >= 10:
                            break

        return experiences

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education from resume."""
        education = []

        edu_section = self._extract_section(text, 'education')
        if not edu_section:
            return education

        # Look for degree patterns
        degree_patterns = [
            (r"(bachelor'?s?)\s*(?:of|in|degree)?\s*([^,\n]{3,50})", 'bachelors'),
            (r"(master'?s?)\s*(?:of|in|degree)?\s*([^,\n]{3,50})", 'masters'),
            (r'(phd|ph\.d\.?|doctorate)\s*(?:of|in)?\s*([^,\n]{3,50})', 'phd'),
            (r'(mba)\s*(?:in)?\s*([^,\n]{0,50})', 'mba'),
            (r'(b\.?s\.?|b\.?a\.?)\s*(?:in)?\s*([^,\n]{3,50})', 'bachelors'),
            (r'(m\.?s\.?|m\.?a\.?)\s*(?:in)?\s*([^,\n]{3,50})', 'masters'),
        ]

        for pattern, degree_type in degree_patterns:
            matches = re.findall(pattern, edu_section, re.IGNORECASE)
            for match in matches:
                degree, field = match if len(match) == 2 else (match, '')
                education.append({
                    'degree': degree.strip(),
                    'degree_type': degree_type,
                    'field': field.strip() if field else '',
                    'institution': '',
                    'graduation_year': ''
                })

        return education[:5]  # Limit to 5 entries

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume."""
        certifications = []

        cert_section = self._extract_section(text, 'certifications')
        if cert_section:
            # Split by common delimiters
            items = re.split(r'[,\n•\|\-]', cert_section)
            for item in items:
                item = item.strip()
                if item and 5 < len(item) < 100:
                    certifications.append(item)

        # Also look for common certification patterns in full text
        cert_patterns = [
            r'(AWS\s+(?:Certified|Solutions|Developer)[^,\n]{0,50})',
            r'(Google\s+Cloud\s+(?:Certified|Professional)[^,\n]{0,50})',
            r'(Azure\s+(?:Certified|Administrator)[^,\n]{0,50})',
            r'(PMP|SCRUM Master|CISSP|CCNA|CCNP)',
            r'(Certified\s+[A-Z][^,\n]{5,50})',
        ]

        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(m.strip() for m in matches)

        # Deduplicate
        seen = set()
        unique = []
        for cert in certifications:
            cert_lower = cert.lower()
            if cert_lower not in seen:
                seen.add(cert_lower)
                unique.append(cert)

        return unique[:20]  # Limit

    def _extract_languages(self, text: str) -> List[str]:
        """Extract languages from resume."""
        languages = []

        # Common languages
        common_languages = [
            'english', 'spanish', 'french', 'german', 'chinese', 'mandarin',
            'japanese', 'korean', 'portuguese', 'italian', 'russian', 'arabic',
            'hindi', 'dutch', 'polish', 'swedish', 'turkish', 'vietnamese'
        ]

        text_lower = text.lower()
        for lang in common_languages:
            if re.search(rf'\b{lang}\b', text_lower):
                languages.append(lang.title())

        return languages

    def _calculate_total_experience(
        self,
        experiences: List[Dict[str, Any]]
    ) -> float:
        """
        Calculate total years of experience.

        Currently uses a heuristic based on number of positions.
        """
        num_jobs = len(experiences)
        if num_jobs == 0:
            return 0.0

        # Rough estimate: average of 2.5 years per position
        # Capped at reasonable maximum
        return min(num_jobs * 2.5, 25.0)
