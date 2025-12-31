"""
Views for AI Matching API

This module provides REST API endpoints for the AI matching service.
"""

import logging
import time
from typing import Optional

from django.shortcuts import get_object_or_404
from django.db import transaction

from rest_framework import status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.pagination import PageNumberPagination

from configurations.models import CandidateProfile, Job
from .models import MatchingResult, RecommendationLog, BiasAuditLog
from .services import (
    MatchingService, RecommendationService, ResumeParserService,
    JobDescriptionAnalyzer, BiasDetectionService
)
from .serializers import (
    MatchCandidatesInputSerializer, MatchJobsInputSerializer,
    ParseResumeInputSerializer, AnalyzeJobDescriptionInputSerializer,
    BiasCheckInputSerializer, CandidateMatchSerializer, JobMatchSerializer,
    ParsedResumeSerializer, JobAnalysisSerializer, BiasReportSerializer,
    MatchingResultSerializer, RecommendationFeedbackSerializer,
    BulkMatchInputSerializer, BulkMatchResultSerializer
)

logger = logging.getLogger(__name__)


# ============================================================================
# Pagination
# ============================================================================

class MatchingResultsPagination(PageNumberPagination):
    """Pagination for matching results."""
    page_size = 20
    page_size_query_param = 'page_size'
    max_page_size = 100


# ============================================================================
# Matching Views
# ============================================================================

class MatchCandidatesView(APIView):
    """
    API endpoint to find matching candidates for a job posting.

    POST /api/ai-matching/match-candidates/
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Find candidates matching a job.

        Request body:
        {
            "job_id": 123,
            "limit": 20,
            "min_score": 0.5,
            "use_ai": true,
            "filters": {"min_experience": 2}
        }

        Returns:
            List of matching candidates with scores
        """
        serializer = MatchCandidatesInputSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )

        data = serializer.validated_data
        start_time = time.time()

        try:
            job = get_object_or_404(Job, id=data['job_id'])

            # Check permission - user should have access to this job
            if not self._can_access_job(request.user, job):
                return Response(
                    {'error': 'You do not have permission to access this job'},
                    status=status.HTTP_403_FORBIDDEN
                )

            # Get candidates
            candidates = CandidateProfile.objects.all()

            # Apply filters
            if data.get('filters'):
                candidates = self._apply_filters(candidates, data['filters'])

            # Calculate matches
            matching_service = MatchingService()
            results = []

            for candidate in candidates[:100]:  # Limit for performance
                try:
                    match_result = matching_service.calculate_match(
                        candidate, job, use_ai=data['use_ai']
                    )

                    if match_result['overall_score'] >= data.get('min_score', 0):
                        results.append({
                            'candidate': candidate,
                            'overall_score': match_result['overall_score'],
                            'skill_score': match_result.get('skill_score', 0),
                            'experience_score': match_result.get('experience_score', 0),
                            'location_score': match_result.get('location_score', 0),
                            'salary_score': match_result.get('salary_score', 0),
                            'matched_skills': match_result.get('matched_skills', []),
                            'missing_skills': match_result.get('missing_skills', []),
                            'confidence': match_result.get('confidence', 'medium'),
                            'algorithm': match_result.get('algorithm', 'unknown')
                        })

                        # Cache the result
                        self._cache_match_result(candidate, job, match_result)

                except Exception as e:
                    logger.warning(f"Error matching candidate {candidate.id}: {e}")

            # Sort by score and limit
            results.sort(key=lambda x: x['overall_score'], reverse=True)
            results = results[:data['limit']]

            processing_time = int((time.time() - start_time) * 1000)

            output_serializer = CandidateMatchSerializer(results, many=True)
            return Response({
                'job_id': job.id,
                'job_title': job.title,
                'total_candidates': len(results),
                'processing_time_ms': processing_time,
                'candidates': output_serializer.data
            })

        except Job.DoesNotExist:
            return Response(
                {'error': 'Job not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            logger.error(f"Error in MatchCandidatesView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _can_access_job(self, user, job) -> bool:
        """Check if user has permission to access job matching."""
        # Admin can access all
        if user.is_staff or user.is_superuser:
            return True

        # Check if user belongs to job's company
        try:
            if hasattr(job.company, 'profile'):
                return job.company.profile.memberships.filter(user=user).exists()
        except Exception:
            pass

        return False

    def _apply_filters(self, queryset, filters):
        """Apply filters to candidate queryset."""
        # Add filter logic as needed
        return queryset

    def _cache_match_result(self, candidate, job, match_result):
        """Cache match result in database."""
        try:
            from django.utils import timezone
            from decimal import Decimal

            MatchingResult.objects.update_or_create(
                candidate=candidate,
                job=job,
                defaults={
                    'overall_score': Decimal(str(match_result['overall_score'])),
                    'skill_score': Decimal(str(match_result.get('skill_score', 0))),
                    'experience_score': Decimal(str(match_result.get('experience_score', 0))),
                    'location_score': Decimal(str(match_result.get('location_score', 0))),
                    'salary_score': Decimal(str(match_result.get('salary_score', 0))),
                    'matching_algorithm': match_result.get('algorithm', 'unknown'),
                    'confidence_level': match_result.get('confidence', 'medium'),
                    'matched_skills': match_result.get('matched_skills', []),
                    'missing_skills': match_result.get('missing_skills', []),
                    'explanation': match_result.get('explanation', {}),
                    'expires_at': timezone.now() + timezone.timedelta(hours=24),
                    'is_stale': False
                }
            )
        except Exception as e:
            logger.warning(f"Failed to cache match result: {e}")


class MatchJobsView(APIView):
    """
    API endpoint to find matching jobs for a candidate.

    POST /api/ai-matching/match-jobs/
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Find jobs matching a candidate's profile.

        Request body:
        {
            "candidate_id": 123,
            "limit": 20,
            "min_score": 0.5,
            "use_ai": true,
            "filters": {"location": "Montreal", "remote_only": true}
        }

        Returns:
            List of matching jobs with scores
        """
        serializer = MatchJobsInputSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )

        data = serializer.validated_data
        start_time = time.time()

        try:
            candidate = get_object_or_404(CandidateProfile, id=data['candidate_id'])

            # Check permission - user should own this profile or be admin
            if not self._can_access_candidate(request.user, candidate):
                return Response(
                    {'error': 'You do not have permission to access this profile'},
                    status=status.HTTP_403_FORBIDDEN
                )

            # Use recommendation service
            recommendation_service = RecommendationService()
            recommendations = recommendation_service.get_jobs_for_candidate(
                candidate,
                limit=data['limit'],
                filters=data.get('filters', {})
            )

            # Filter by min_score
            min_score = data.get('min_score', 0)
            results = []
            for rec in recommendations:
                if rec['score'] >= min_score:
                    match_details = rec.get('match_details', {})
                    results.append({
                        'job': rec['job'],
                        'overall_score': rec['score'],
                        'skill_score': match_details.get('skill_score', 0),
                        'experience_score': match_details.get('experience_score', 0),
                        'location_score': match_details.get('location_score', 0),
                        'salary_score': match_details.get('salary_score', 0),
                        'matched_skills': match_details.get('matched_skills', []),
                        'missing_skills': match_details.get('missing_skills', []),
                        'confidence': match_details.get('confidence', 'medium'),
                        'algorithm': match_details.get('algorithm', 'unknown')
                    })

            processing_time = int((time.time() - start_time) * 1000)

            output_serializer = JobMatchSerializer(results, many=True)
            return Response({
                'candidate_id': candidate.id,
                'total_jobs': len(results),
                'processing_time_ms': processing_time,
                'jobs': output_serializer.data
            })

        except CandidateProfile.DoesNotExist:
            return Response(
                {'error': 'Candidate profile not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            logger.error(f"Error in MatchJobsView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _can_access_candidate(self, user, candidate) -> bool:
        """Check if user has permission to access candidate matching."""
        if user.is_staff or user.is_superuser:
            return True
        return candidate.user == user


# ============================================================================
# Resume Parsing View
# ============================================================================

class ParseResumeView(APIView):
    """
    API endpoint to parse a resume and extract structured information.

    POST /api/ai-matching/parse-resume/
    """
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    def post(self, request):
        """
        Parse resume and extract skills, experience, education.

        Request body:
        {
            "resume_text": "Full resume text...",
            "candidate_id": 123  // Optional
        }
        OR multipart form with resume_file

        Returns:
            Parsed resume data
        """
        serializer = ParseResumeInputSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )

        data = serializer.validated_data
        start_time = time.time()

        try:
            # Get resume text
            resume_text = data.get('resume_text', '')

            if not resume_text and data.get('resume_file'):
                resume_text = self._extract_text_from_file(data['resume_file'])

            if not resume_text:
                return Response(
                    {'error': 'Could not extract text from resume'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Parse resume
            parser = ResumeParserService()
            parsed = parser.execute(resume_text)

            processing_time = int((time.time() - start_time) * 1000)

            # Optionally update candidate profile
            if data.get('candidate_id'):
                self._update_candidate_profile(data['candidate_id'], parsed, request.user)

            output_serializer = ParsedResumeSerializer({
                'skills': parsed.skills,
                'experience_years': parsed.experience_years,
                'education': parsed.education,
                'work_history': parsed.work_history,
                'certifications': parsed.certifications,
                'summary': parsed.summary
            })

            return Response({
                'processing_time_ms': processing_time,
                'parsed_resume': output_serializer.data
            })

        except Exception as e:
            logger.error(f"Error in ParseResumeView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _extract_text_from_file(self, file) -> str:
        """Extract text from uploaded file."""
        filename = file.name.lower()

        if filename.endswith('.txt'):
            return file.read().decode('utf-8', errors='ignore')

        elif filename.endswith('.pdf'):
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() or ''
                return text
            except ImportError:
                logger.warning("PyPDF2 not installed for PDF parsing")
            except Exception as e:
                logger.warning(f"PDF parsing failed: {e}")

        elif filename.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(file)
                return '\n'.join([para.text for para in doc.paragraphs])
            except ImportError:
                logger.warning("python-docx not installed for DOCX parsing")
            except Exception as e:
                logger.warning(f"DOCX parsing failed: {e}")

        return ''

    def _update_candidate_profile(self, candidate_id, parsed, user):
        """Update candidate profile with parsed data."""
        try:
            candidate = CandidateProfile.objects.get(id=candidate_id)

            # Verify permission
            if candidate.user != user and not user.is_staff:
                return

            # Update bio if empty
            if not candidate.bio and parsed.summary:
                candidate.bio = parsed.summary
                candidate.save()

            # Add extracted skills
            from configurations.models import Skill
            for skill_name in parsed.skills:
                skill, _ = Skill.objects.get_or_create(
                    name=skill_name.title(),
                    defaults={'description': ''}
                )
                candidate.skills.add(skill)

        except CandidateProfile.DoesNotExist:
            pass
        except Exception as e:
            logger.warning(f"Failed to update candidate profile: {e}")


# ============================================================================
# Job Description Analysis View
# ============================================================================

class AnalyzeJobDescriptionView(APIView):
    """
    API endpoint to analyze a job description.

    POST /api/ai-matching/analyze-job/
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Analyze job description and extract requirements.

        Request body:
        {
            "job_description": "Full job description...",
            "job_title": "Software Engineer",
            "job_id": 123  // Optional
        }

        Returns:
            Analyzed job requirements
        """
        serializer = AnalyzeJobDescriptionInputSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )

        data = serializer.validated_data
        start_time = time.time()

        try:
            analyzer = JobDescriptionAnalyzer()
            analysis = analyzer.execute(
                data['job_description'],
                data.get('job_title', '')
            )

            processing_time = int((time.time() - start_time) * 1000)

            # Optionally update job embedding
            if data.get('job_id'):
                self._update_job_embedding(data['job_id'], analysis)

            output_serializer = JobAnalysisSerializer({
                'required_skills': analysis.required_skills,
                'preferred_skills': analysis.preferred_skills,
                'experience_range': list(analysis.experience_range),
                'education_level': analysis.education_level,
                'is_remote': analysis.is_remote,
                'salary_range': list(analysis.salary_range) if analysis.salary_range else None,
                'key_responsibilities': analysis.key_responsibilities,
                'company_values': analysis.company_values
            })

            return Response({
                'processing_time_ms': processing_time,
                'analysis': output_serializer.data
            })

        except Exception as e:
            logger.error(f"Error in AnalyzeJobDescriptionView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _update_job_embedding(self, job_id, analysis):
        """Update job embedding with analysis data."""
        try:
            from .models import JobEmbedding

            job = Job.objects.get(id=job_id)
            embedding, _ = JobEmbedding.objects.get_or_create(
                job=job,
                defaults={'embedding_vector': [0.0] * 1536}
            )

            embedding.skills_extracted = (
                analysis.required_skills + analysis.preferred_skills
            )
            embedding.experience_years_min = analysis.experience_range[0]
            embedding.experience_years_max = analysis.experience_range[1]
            embedding.is_remote = analysis.is_remote
            embedding.save()

        except Job.DoesNotExist:
            pass
        except Exception as e:
            logger.warning(f"Failed to update job embedding: {e}")


# ============================================================================
# Bias Detection View
# ============================================================================

class BiasCheckView(APIView):
    """
    API endpoint to check text for potential bias.

    POST /api/ai-matching/bias-check/
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Check text for bias (gender, age, etc.).

        Request body:
        {
            "text": "Job description text...",
            "content_type": "job_posting",
            "content_id": 123,  // Optional
            "log_audit": true
        }

        Returns:
            Bias detection report
        """
        serializer = BiasCheckInputSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )

        data = serializer.validated_data
        start_time = time.time()

        try:
            detector = BiasDetectionService()
            report = detector.execute(data['text'], data['content_type'])

            processing_time = int((time.time() - start_time) * 1000)

            # Log audit if requested
            if data.get('log_audit', True) and data.get('content_id'):
                detector.log_audit(
                    content_id=data['content_id'],
                    content_type=data['content_type'],
                    report=report,
                    user=request.user
                )

            output_serializer = BiasReportSerializer({
                'has_bias': report.has_bias,
                'bias_score': report.bias_score,
                'gender_bias': report.gender_bias,
                'age_bias': report.age_bias,
                'other_bias': report.other_bias,
                'suggestions': report.suggestions
            })

            return Response({
                'processing_time_ms': processing_time,
                'report': output_serializer.data
            })

        except Exception as e:
            logger.error(f"Error in BiasCheckView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


# ============================================================================
# Recommendation Feedback View
# ============================================================================

class RecommendationFeedbackView(APIView):
    """
    API endpoint to submit feedback on recommendations.

    POST /api/ai-matching/recommendation-feedback/
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Submit feedback on recommendation quality.

        Request body:
        {
            "recommendation_id": "uuid-here",
            "items_viewed": [1, 2, 3],
            "items_clicked": [1],
            "items_applied": [],
            "rating": 4,
            "feedback": "Good recommendations!"
        }
        """
        serializer = RecommendationFeedbackSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )

        data = serializer.validated_data

        try:
            log = get_object_or_404(
                RecommendationLog,
                uuid=data['recommendation_id']
            )

            # Verify ownership
            if log.user != request.user and not request.user.is_staff:
                return Response(
                    {'error': 'You do not have permission to update this recommendation'},
                    status=status.HTTP_403_FORBIDDEN
                )

            # Update log with feedback
            if 'items_viewed' in data:
                log.items_viewed = list(set(log.items_viewed + data['items_viewed']))
            if 'items_clicked' in data:
                log.items_clicked = list(set(log.items_clicked + data['items_clicked']))
            if 'items_applied' in data:
                log.items_applied = list(set(log.items_applied + data['items_applied']))
            if 'rating' in data:
                log.user_rating = data['rating']
            if 'feedback' in data:
                log.user_feedback = data['feedback']

            log.save()

            return Response({'status': 'Feedback recorded'})

        except RecommendationLog.DoesNotExist:
            return Response(
                {'error': 'Recommendation not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            logger.error(f"Error in RecommendationFeedbackView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


# ============================================================================
# Bulk Match View (Admin Only)
# ============================================================================

class BulkMatchView(APIView):
    """
    API endpoint for bulk matching operations.

    POST /api/ai-matching/bulk-match/
    """
    permission_classes = [IsAdminUser]

    def post(self, request):
        """
        Trigger bulk matching calculation.

        Request body:
        {
            "job_ids": [1, 2, 3],  // Optional
            "candidate_ids": [1, 2, 3],  // Optional
            "recalculate": false
        }
        """
        serializer = BulkMatchInputSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )

        data = serializer.validated_data
        start_time = time.time()

        try:
            # Import tasks
            from .tasks import calculate_match_scores

            # Trigger async task
            task = calculate_match_scores.delay(
                job_ids=data.get('job_ids'),
                candidate_ids=data.get('candidate_ids'),
                recalculate=data.get('recalculate', False)
            )

            return Response({
                'status': 'Bulk matching task queued',
                'task_id': str(task.id) if hasattr(task, 'id') else None
            }, status=status.HTTP_202_ACCEPTED)

        except Exception as e:
            logger.error(f"Error in BulkMatchView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


# ============================================================================
# Match Results History View
# ============================================================================

class MatchResultsHistoryView(APIView):
    """
    API endpoint to get historical match results.

    GET /api/ai-matching/match-results/
    """
    permission_classes = [IsAuthenticated]
    pagination_class = MatchingResultsPagination

    def get(self, request):
        """
        Get historical match results for current user.

        Query params:
        - job_id: Filter by job
        - candidate_id: Filter by candidate
        - min_score: Minimum score threshold
        """
        try:
            queryset = MatchingResult.objects.all()

            # Filter by user access
            if not request.user.is_staff:
                # Get candidate profiles owned by user
                candidate_ids = CandidateProfile.objects.filter(
                    user=request.user
                ).values_list('id', flat=True)
                queryset = queryset.filter(candidate_id__in=candidate_ids)

            # Apply filters
            job_id = request.query_params.get('job_id')
            if job_id:
                queryset = queryset.filter(job_id=job_id)

            candidate_id = request.query_params.get('candidate_id')
            if candidate_id:
                queryset = queryset.filter(candidate_id=candidate_id)

            min_score = request.query_params.get('min_score')
            if min_score:
                queryset = queryset.filter(overall_score__gte=float(min_score))

            # Paginate
            paginator = self.pagination_class()
            page = paginator.paginate_queryset(queryset, request)

            if page is not None:
                serializer = MatchingResultSerializer(page, many=True)
                return paginator.get_paginated_response(serializer.data)

            serializer = MatchingResultSerializer(queryset, many=True)
            return Response(serializer.data)

        except Exception as e:
            logger.error(f"Error in MatchResultsHistoryView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


# ============================================================================
# AI Service Health View
# ============================================================================

class AIServiceHealthView(APIView):
    """
    API endpoint to check AI service health status.

    GET /api/ai-matching/health/
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        """Get AI service health status."""
        from .models import AIServiceStatus

        try:
            statuses = AIServiceStatus.objects.all()
            services = {}

            for status_obj in statuses:
                services[status_obj.service_name] = {
                    'is_available': status_obj.is_available,
                    'last_check': status_obj.last_check.isoformat() if status_obj.last_check else None,
                    'failure_count': status_obj.failure_count,
                    'requests_today': status_obj.requests_today,
                    'daily_limit': status_obj.daily_limit
                }

            # Check if any critical service is down
            overall_healthy = all(
                s.get('is_available', True)
                for s in services.values()
            )

            return Response({
                'healthy': overall_healthy,
                'services': services,
                'fallback_available': True  # Rule-based fallback always available
            })

        except Exception as e:
            logger.error(f"Error in AIServiceHealthView: {e}")
            return Response({
                'healthy': False,
                'error': str(e),
                'fallback_available': True
            })


# ============================================================================
# Cycle 7 - Enhanced API Views with Caching and Permissions
# ============================================================================

from django.core.cache import cache
from django.utils.decorators import method_decorator
from django.views.decorators.cache import cache_page

from .serializers import (
    MatchResultSerializer, JobMatchResultSerializer,
    MatchExplanationSerializer, BiasReportDetailSerializer,
    ResumeParseResultSerializer, MatchingProfileSerializer,
    CandidateMinimalSerializer, JobListSerializer
)


class CandidateMatchingView(APIView):
    """
    Enhanced candidate matching view with caching support.

    GET /api/ai-matching/candidates/match/
        - Get top matches for a job (cached)

    POST /api/ai-matching/candidates/match/
        - Compute fresh matches (bypasses cache)
    """
    permission_classes = [IsAuthenticated]

    CACHE_TIMEOUT = 300  # 5 minutes

    def get(self, request):
        """
        Get top matching candidates for a job.

        Query params:
        - job_id: Required job ID
        - limit: Max results (default 20, max 100)
        - min_score: Minimum match score (0-1)
        - use_cache: Use cached results (default true)
        """
        job_id = request.query_params.get('job_id')
        if not job_id:
            return Response(
                {'error': 'job_id is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        limit = min(int(request.query_params.get('limit', 20)), 100)
        min_score = float(request.query_params.get('min_score', 0))
        use_cache = request.query_params.get('use_cache', 'true').lower() == 'true'

        # Check cache
        cache_key = f"candidate_matches_{job_id}_{limit}_{min_score}"
        if use_cache:
            cached_result = cache.get(cache_key)
            if cached_result:
                return Response(cached_result)

        try:
            job = get_object_or_404(Job, id=job_id)

            # Permission check
            if not request.user.is_staff:
                if hasattr(job, 'company') and hasattr(job.company, 'profile'):
                    if not job.company.profile.memberships.filter(user=request.user).exists():
                        return Response(
                            {'error': 'Permission denied'},
                            status=status.HTTP_403_FORBIDDEN
                        )

            # Get cached matching results from database
            cached_matches = MatchingResult.objects.filter(
                job=job,
                is_stale=False,
                overall_score__gte=min_score
            ).select_related('candidate__user').order_by('-overall_score')[:limit]

            results = []
            for match in cached_matches:
                candidate = match.candidate
                results.append({
                    'candidate': {
                        'id': candidate.id,
                        'uuid': str(candidate.uuid) if hasattr(candidate, 'uuid') else None,
                        'email': candidate.user.email,
                        'full_name': f"{candidate.user.first_name} {candidate.user.last_name}".strip() or candidate.user.username,
                        'avatar_url': None,
                        'title': getattr(candidate, 'title', ''),
                        'location': getattr(candidate, 'city', '') or '',
                    },
                    'overall_score': float(match.overall_score),
                    'skill_match_score': float(match.skill_score or 0),
                    'experience_match_score': float(match.experience_score or 0),
                    'location_score': float(match.location_score or 1),
                    'salary_score': float(match.salary_score or 1),
                    'culture_score': float(match.culture_score or 1),
                    'explanation': match.explanation.get('reasons', []) if match.explanation else [],
                    'matched_skills': match.matched_skills,
                    'missing_skills': match.missing_skills,
                    'confidence': match.confidence_level,
                    'algorithm': match.matching_algorithm,
                })

            response_data = {
                'job_id': job.id,
                'job_title': job.title,
                'total_matches': len(results),
                'matches': results,
                'cached': use_cache and len(results) > 0,
            }

            # Cache the result
            if results:
                cache.set(cache_key, response_data, self.CACHE_TIMEOUT)

            return Response(response_data)

        except Exception as e:
            logger.error(f"Error in CandidateMatchingView.get: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def post(self, request):
        """
        Compute fresh matches for a job.
        Triggers async computation and returns task ID.
        """
        job_id = request.data.get('job_id')
        if not job_id:
            return Response(
                {'error': 'job_id is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            job = get_object_or_404(Job, id=job_id)

            # Permission check
            if not request.user.is_staff:
                return Response(
                    {'error': 'Admin access required'},
                    status=status.HTTP_403_FORBIDDEN
                )

            # Invalidate cache
            cache.delete_pattern(f"candidate_matches_{job_id}_*")

            # Trigger async computation
            from .tasks import calculate_match_scores
            task = calculate_match_scores.delay(
                job_ids=[job_id],
                recalculate=True
            )

            return Response({
                'status': 'Computing matches',
                'task_id': str(task.id) if hasattr(task, 'id') else None,
                'job_id': job_id,
            }, status=status.HTTP_202_ACCEPTED)

        except Exception as e:
            logger.error(f"Error in CandidateMatchingView.post: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class JobMatchingView(APIView):
    """
    Match jobs to a candidate's profile.

    GET /api/ai-matching/jobs/match/
    """
    permission_classes = [IsAuthenticated]

    CACHE_TIMEOUT = 300

    def get(self, request):
        """
        Get matching jobs for a candidate.

        Query params:
        - candidate_id: Required candidate profile ID
        - limit: Max results (default 20)
        - min_score: Minimum match score (0-1)
        - remote_only: Filter to remote jobs only
        - location: Filter by location
        """
        candidate_id = request.query_params.get('candidate_id')

        # If no candidate_id, try to get current user's profile
        if not candidate_id:
            try:
                candidate = CandidateProfile.objects.get(user=request.user)
                candidate_id = candidate.id
            except CandidateProfile.DoesNotExist:
                return Response(
                    {'error': 'candidate_id is required or you must have a candidate profile'},
                    status=status.HTTP_400_BAD_REQUEST
                )

        limit = min(int(request.query_params.get('limit', 20)), 100)
        min_score = float(request.query_params.get('min_score', 0))

        try:
            candidate = get_object_or_404(CandidateProfile, id=candidate_id)

            # Permission check
            if not request.user.is_staff and candidate.user != request.user:
                return Response(
                    {'error': 'Permission denied'},
                    status=status.HTTP_403_FORBIDDEN
                )

            # Check cache
            cache_key = f"job_matches_{candidate_id}_{limit}_{min_score}"
            cached_result = cache.get(cache_key)
            if cached_result:
                return Response(cached_result)

            # Get cached matching results
            cached_matches = MatchingResult.objects.filter(
                candidate=candidate,
                is_stale=False,
                overall_score__gte=min_score
            ).select_related('job').order_by('-overall_score')[:limit]

            results = []
            for match in cached_matches:
                job = match.job
                company_name = "Unknown"
                try:
                    company_name = job.company.name if hasattr(job, 'company') else "Unknown"
                except Exception:
                    pass

                why_good_fit = []
                if match.matched_skills:
                    why_good_fit.append(f"Matches {len(match.matched_skills)} of your skills")
                if float(match.experience_score or 0) >= 0.8:
                    why_good_fit.append("Your experience level is a great fit")
                if float(match.overall_score) >= 0.85:
                    why_good_fit.append("Strong overall alignment with your profile")

                results.append({
                    'job': {
                        'id': job.id,
                        'uuid': str(job.uuid) if hasattr(job, 'uuid') else None,
                        'title': job.title,
                        'company_name': company_name,
                        'location': getattr(job, 'location', ''),
                        'is_remote': getattr(job, 'is_remote', False),
                        'salary_min': float(job.salary_from) if hasattr(job, 'salary_from') and job.salary_from else None,
                        'salary_max': float(job.salary_to) if hasattr(job, 'salary_to') and job.salary_to else None,
                        'posted_at': job.created_at.isoformat() if hasattr(job, 'created_at') else None,
                    },
                    'overall_score': float(match.overall_score),
                    'skill_match_score': float(match.skill_score or 0),
                    'experience_match_score': float(match.experience_score or 0),
                    'why_good_fit': why_good_fit,
                    'growth_opportunities': match.missing_skills[:5] if match.missing_skills else [],
                })

            response_data = {
                'candidate_id': candidate.id,
                'total_matches': len(results),
                'matches': results,
            }

            if results:
                cache.set(cache_key, response_data, self.CACHE_TIMEOUT)

            return Response(response_data)

        except Exception as e:
            logger.error(f"Error in JobMatchingView.get: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class MatchExplanationView(APIView):
    """
    Get detailed explanation for a specific match.

    GET /api/ai-matching/match/<uuid:match_id>/explain/
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, match_id):
        """
        Get human-readable explanation of match.
        """
        try:
            match = get_object_or_404(MatchingResult, uuid=match_id)

            # Permission check
            if not request.user.is_staff:
                if match.candidate.user != request.user:
                    # Check if user has access to the job
                    has_job_access = False
                    if hasattr(match.job, 'company') and hasattr(match.job.company, 'profile'):
                        has_job_access = match.job.company.profile.memberships.filter(
                            user=request.user
                        ).exists()

                    if not has_job_access:
                        return Response(
                            {'error': 'Permission denied'},
                            status=status.HTTP_403_FORBIDDEN
                        )

            # Build detailed explanation
            explanation_data = {
                'match_id': str(match.uuid),
                'candidate_id': match.candidate.id,
                'job_id': match.job.id,
                'overall_score': float(match.overall_score),
                'score_breakdown': {
                    'skill_score': float(match.skill_score or 0),
                    'experience_score': float(match.experience_score or 0),
                    'location_score': float(match.location_score or 0),
                    'salary_score': float(match.salary_score or 0),
                    'culture_score': float(match.culture_score or 0),
                    'education_score': float(match.education_score or 0),
                },
                'skill_analysis': {
                    'matched': match.matched_skills,
                    'missing': match.missing_skills,
                    'match_percentage': round(
                        len(match.matched_skills) / max(len(match.matched_skills) + len(match.missing_skills), 1) * 100,
                        1
                    ),
                },
                'experience_analysis': match.explanation.get('experience_analysis', {}),
                'location_analysis': match.explanation.get('location_analysis', {}),
                'salary_analysis': match.explanation.get('salary_analysis', {}),
                'human_readable_summary': self._generate_summary(match),
                'improvement_suggestions': self._generate_suggestions(match),
                'algorithm_details': {
                    'algorithm': match.matching_algorithm,
                    'confidence': match.confidence_level,
                    'calculated_at': match.calculated_at.isoformat(),
                    'expires_at': match.expires_at.isoformat() if match.expires_at else None,
                },
            }

            return Response(explanation_data)

        except MatchingResult.DoesNotExist:
            return Response(
                {'error': 'Match not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            logger.error(f"Error in MatchExplanationView: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _generate_summary(self, match):
        """Generate human-readable summary."""
        score = float(match.overall_score)
        if score >= 0.9:
            quality = "excellent"
        elif score >= 0.75:
            quality = "strong"
        elif score >= 0.6:
            quality = "good"
        elif score >= 0.4:
            quality = "moderate"
        else:
            quality = "limited"

        matched_count = len(match.matched_skills)
        missing_count = len(match.missing_skills)

        summary = f"This is a {quality} match with an overall score of {score:.0%}. "
        summary += f"The candidate matches {matched_count} of the required skills"
        if missing_count > 0:
            summary += f" and is missing {missing_count} skills"
        summary += "."

        return summary

    def _generate_suggestions(self, match):
        """Generate improvement suggestions."""
        suggestions = []

        if match.missing_skills:
            top_missing = match.missing_skills[:3]
            suggestions.append(
                f"Developing skills in {', '.join(top_missing)} would improve this match."
            )

        if float(match.experience_score or 0) < 0.6:
            suggestions.append(
                "Gaining more relevant experience would strengthen this candidacy."
            )

        if float(match.skill_score or 0) < 0.6:
            suggestions.append(
                "Highlighting transferable skills could improve the skill match score."
            )

        return suggestions


class BiasDetectionView(APIView):
    """
    Bias detection and fairness analysis for job postings.

    GET /api/ai-matching/bias-report/
    POST /api/ai-matching/bias-report/
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        """
        Get bias report for a job posting.

        Query params:
        - job_id: Job posting ID
        - include_recommendations: Include fix suggestions (default true)
        """
        job_id = request.query_params.get('job_id')
        if not job_id:
            return Response(
                {'error': 'job_id is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            job = get_object_or_404(Job, id=job_id)

            # Get most recent bias audit
            audit = BiasAuditLog.objects.filter(
                content_type='job_posting',
                content_id=job_id
            ).order_by('-created_at').first()

            if not audit:
                return Response({
                    'message': 'No bias analysis found for this job. Use POST to run analysis.',
                    'job_id': job_id,
                })

            return Response({
                'job_id': job_id,
                'job_title': job.title,
                'overall_fairness_score': 1 - float(audit.bias_score or 0),
                'has_bias': audit.bias_detected,
                'demographic_analysis': {
                    'gender_bias': audit.bias_types and 'gender' in audit.bias_types,
                    'age_bias': audit.bias_types and 'age' in audit.bias_types,
                    'other_bias': [b for b in (audit.bias_types or []) if b not in ['gender', 'age']],
                },
                'flagged_phrases': audit.flagged_phrases,
                'recommendations': audit.suggestions,
                'severity': self._calculate_severity(audit.bias_score),
                'audit_id': str(audit.uuid),
                'analyzed_at': audit.created_at.isoformat(),
            })

        except Exception as e:
            logger.error(f"Error in BiasDetectionView.get: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def post(self, request):
        """
        Run bias detection on text or job posting.
        """
        job_id = request.data.get('job_id')
        text = request.data.get('text')

        if not job_id and not text:
            return Response(
                {'error': 'Either job_id or text is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            if job_id:
                job = get_object_or_404(Job, id=job_id)
                text = f"{job.title}\n\n{job.description}"

            # Run bias detection
            detector = BiasDetectionService()
            report = detector.execute(text, 'job_posting')

            # Log the audit
            content_id = job_id if job_id else 0
            audit = BiasAuditLog.objects.create(
                content_type='job_posting',
                content_id=content_id,
                bias_detected=report.has_bias,
                bias_types=list(set(
                    (['gender'] if report.gender_bias else []) +
                    (['age'] if report.age_bias else [])
                )),
                bias_score=report.bias_score,
                flagged_phrases=report.suggestions,
                suggestions=report.suggestions,
                auditor=request.user,
                automated=True,
            )

            return Response({
                'overall_fairness_score': 1 - report.bias_score,
                'has_bias': report.has_bias,
                'demographic_analysis': {
                    'gender': report.gender_bias,
                    'age_bias': report.age_bias,
                    'other_bias': report.other_bias,
                },
                'flagged_phrases': [],
                'recommendations': report.suggestions,
                'severity': self._calculate_severity(report.bias_score),
                'audit_id': str(audit.uuid),
            })

        except Exception as e:
            logger.error(f"Error in BiasDetectionView.post: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _calculate_severity(self, bias_score):
        """Calculate severity level from bias score."""
        if not bias_score:
            return 'none'
        score = float(bias_score)
        if score >= 0.8:
            return 'critical'
        elif score >= 0.6:
            return 'high'
        elif score >= 0.4:
            return 'medium'
        elif score >= 0.2:
            return 'low'
        return 'none'


class ResumeParsingView(APIView):
    """
    Parse uploaded resume and return structured data.

    POST /api/ai-matching/resume/parse/
    """
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    def post(self, request):
        """
        Parse resume and optionally update candidate profile.

        Request:
        - resume_file: File upload (PDF, DOCX, TXT)
        - resume_text: Or plain text
        - candidate_id: Optional - update this profile
        - create_profile: If true and no candidate_id, create new profile
        """
        start_time = time.time()

        # Get resume content
        resume_text = request.data.get('resume_text', '')
        resume_file = request.FILES.get('resume_file')
        candidate_id = request.data.get('candidate_id')
        create_profile = request.data.get('create_profile', 'false').lower() == 'true'

        if not resume_text and not resume_file:
            return Response(
                {'error': 'Either resume_text or resume_file is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            # Extract text from file if needed
            file_info = None
            if resume_file and not resume_text:
                resume_text = self._extract_text_from_file(resume_file)
                file_info = {
                    'name': resume_file.name,
                    'size': resume_file.size,
                    'content_type': resume_file.content_type,
                }

            if not resume_text:
                return Response(
                    {'error': 'Could not extract text from resume'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Parse resume
            parser = ResumeParserService()
            parsed = parser.execute(resume_text)

            processing_time = int((time.time() - start_time) * 1000)

            # Optionally update/create profile
            updated_candidate_id = None
            if candidate_id:
                self._update_candidate_profile(candidate_id, parsed, request.user)
                updated_candidate_id = candidate_id
            elif create_profile:
                updated_candidate_id = self._create_candidate_profile(parsed, request.user)

            return Response({
                'success': True,
                'candidate_id': updated_candidate_id,
                'parsed_data': {
                    'skills': parsed.skills,
                    'experience_years': parsed.experience_years,
                    'education': parsed.education,
                    'work_history': parsed.work_history,
                    'certifications': parsed.certifications,
                    'summary': parsed.summary,
                },
                'confidence_scores': {
                    'skills': 0.85,
                    'experience': 0.80,
                    'education': 0.90,
                },
                'warnings': [],
                'processing_time_ms': processing_time,
                'file_info': file_info,
            })

        except Exception as e:
            logger.error(f"Error in ResumeParsingView: {e}")
            return Response({
                'success': False,
                'error': str(e),
                'processing_time_ms': int((time.time() - start_time) * 1000),
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def _extract_text_from_file(self, file) -> str:
        """Extract text from uploaded file."""
        filename = file.name.lower()

        if filename.endswith('.txt'):
            return file.read().decode('utf-8', errors='ignore')

        elif filename.endswith('.pdf'):
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() or ''
                return text
            except Exception as e:
                logger.warning(f"PDF parsing failed: {e}")

        elif filename.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(file)
                return '\n'.join([para.text for para in doc.paragraphs])
            except Exception as e:
                logger.warning(f"DOCX parsing failed: {e}")

        return ''

    def _update_candidate_profile(self, candidate_id, parsed, user):
        """Update existing candidate profile."""
        try:
            candidate = CandidateProfile.objects.get(id=candidate_id)
            if candidate.user != user and not user.is_staff:
                return

            if not candidate.bio and parsed.summary:
                candidate.bio = parsed.summary
                candidate.save()

            from configurations.models import Skill
            for skill_name in parsed.skills:
                skill, _ = Skill.objects.get_or_create(
                    name=skill_name.title(),
                    defaults={'description': ''}
                )
                candidate.skills.add(skill)

        except Exception as e:
            logger.warning(f"Failed to update candidate profile: {e}")

    def _create_candidate_profile(self, parsed, user):
        """Create new candidate profile."""
        try:
            candidate, created = CandidateProfile.objects.get_or_create(
                user=user,
                defaults={'bio': parsed.summary or ''}
            )

            from configurations.models import Skill
            for skill_name in parsed.skills:
                skill, _ = Skill.objects.get_or_create(
                    name=skill_name.title(),
                    defaults={'description': ''}
                )
                candidate.skills.add(skill)

            return candidate.id
        except Exception as e:
            logger.warning(f"Failed to create candidate profile: {e}")
            return None
